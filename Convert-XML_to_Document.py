import os
import xml.etree.ElementTree as ET
from docx import Document

# Function to process each XML file
def process_xml_file(xml_file_path, output_doc_path):
    tree = ET.parse(xml_file_path)
    root = tree.getroot()

    # List to store pairs of (text, timecode)
    entries = []

    # Temp variables to store current text and TC
    current_text = None
    current_tc = None

    # Iterate through the XML to find the required values
    for list_elem in root.findall(".//ListElem"):
        # Create a dictionary to store the values found in AvProp elements within this ListElem
        avprop_values = {}

        # Loop through each AvProp element inside ListElem and store the values in the dictionary
        for avprop in list_elem.findall("AvProp"):
            name_attr = avprop.get('name')
            if name_attr and avprop.text:
                avprop_values[name_attr] = avprop.text.strip()

        # Check if the element contains the text
        if avprop_values.get('OMFI:ATTB:Name') == '_ATN_CRM_COM':
            current_text = avprop_values.get('OMFI:ATTB:StringAttribute')
            #print(f"Found Text: {current_text}")

        # Check if the element contains the timecode
        if avprop_values.get('OMFI:ATTB:Name') == '_ATN_CRM_TC':
            current_tc = avprop_values.get('OMFI:ATTB:StringAttribute')
            #print(f"Found TC: {current_tc}")

        # When both text and TC are found, add them as a pair to the entries
        if current_text and current_tc:
            entries.append((current_tc, current_text))
            current_text = None  # Reset text for the next pair
            current_tc = None    # Reset TC for the next pair

    # Create Word document
    doc = Document()

    # Function to process the text
    def clean_text(text):
        # Replace "/13" or "/13/13" with newlines, except if it's at the end
        if text.endswith("/13") or text.endswith("/13/13"):
            text = text.rsplit("/13", 1)[0]  # Remove the ending /13 or /13/13
        # Replace any occurrences of "/13" or "/13/13" within the text with a new line
        text = text.replace("/13/13", "\n").replace("/13", "\n")
        return text

    # Write sorted timecodes and text to the Word document
    for tc, text in entries:
        doc.add_paragraph(tc, style='Normal')  # Add TC in one paragraph
        cleaned_text = clean_text(text)  # Clean the text before adding it
        doc.add_paragraph(cleaned_text, style='Normal')  # Add the cleaned text
        doc.add_paragraph()  # Add an empty paragraph for spacing between entries

    # Save the document
    doc.save(output_doc_path)
    print(f"Document saved: {output_doc_path}")

# Define the input folder path
input_folder = "XML"

# Create the input folder if it doesn't exist
os.makedirs(input_folder, exist_ok=True)

# Start processing XML files
for file_name in os.listdir(input_folder):
    if file_name.lower().endswith('.xml'):
        xml_file_path = os.path.join(input_folder, file_name)
        output_doc_path = os.path.join(input_folder, f"{os.path.splitext(file_name)[0]}.docx")

        # Check if the output Word document already exists
        if not os.path.exists(output_doc_path):
            process_xml_file(xml_file_path, output_doc_path)
        else:
            print(f"Document '{output_doc_path}' already exists. Skipping...")